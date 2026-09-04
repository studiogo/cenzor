#!/usr/bin/env python3
"""Zamienia dokument na zwykly tekst — PDF, DOCX, ODT albo plik tekstowy.

Sam rozpoznaje format po zawartosci, nie po rozszerzeniu, bo ludzie
zapisuja PDF-y pod nazwa .doc i odwrotnie.
"""
import sys, zipfile
from pathlib import Path


class BladOdczytu(Exception):
    pass


def _pdf(sciezka):
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        raise BladOdczytu("Brak biblioteki do PDF. Uruchom: pip install pdfminer.six")
    tekst = extract_text(str(sciezka)) or ""
    if len(tekst.split()) < 5:
        raise BladOdczytu(
            f"{sciezka.name} to skan — sam obraz, bez warstwy tekstowej.\n"
            "Zadne narzedzie nie wytnie danych z obrazka. Przepusc plik przez\n"
            "rozpoznawanie tekstu (OCR) albo popros o wersje tekstowa.")
    return tekst


def _docx(sciezka):
    try:
        import docx
    except ImportError:
        raise BladOdczytu("Brak biblioteki do Worda. Uruchom: pip install python-docx")
    d = docx.Document(str(sciezka))
    czesci = [a.text for a in d.paragraphs]
    for t in d.tables:                      # tabele leza poza akapitami
        for wiersz in t.rows:
            czesci.append("\t".join(k.text for k in wiersz.cells))
    for sekcja in d.sections:               # naglowki i stopki tez miewaja dane
        for obszar in (sekcja.header, sekcja.footer):
            czesci += [a.text for a in obszar.paragraphs]
    return "\n".join(czesci)


def _odt(sciezka):
    import re, xml.etree.ElementTree as ET
    with zipfile.ZipFile(sciezka) as z:
        xml = z.read("content.xml").decode("utf-8", "replace")
    tekst = re.sub(r"<[^>]+>", "\n", xml)
    return re.sub(r"\n{3,}", "\n\n", tekst)


def _rodzaj(sciezka):
    naglowek = sciezka.open("rb").read(8)
    if naglowek.startswith(b"%PDF"):
        return "pdf"
    if naglowek.startswith(b"PK"):          # DOCX i ODT to spakowane katalogi
        try:
            with zipfile.ZipFile(sciezka) as z:
                nazwy = z.namelist()
            if any(n.startswith("word/") for n in nazwy):
                return "docx"
            if "content.xml" in nazwy:
                return "odt"
        except zipfile.BadZipFile:
            pass
        raise BladOdczytu(f"{sciezka.name}: nieznany format spakowany.")
    if naglowek.startswith(b"\xd0\xcf\x11\xe0"):
        raise BladOdczytu(
            f"{sciezka.name} to stary Word (.doc sprzed 2007).\n"
            "Otworz go i zapisz jako .docx albo .txt.")
    return "txt"


def wczytaj(sciezka):
    """Zwraca (tekst, rozpoznany_format). Rzuca BladOdczytu z rada, co zrobic."""
    sciezka = Path(sciezka)
    if not sciezka.exists():
        raise BladOdczytu(f"Nie ma pliku: {sciezka}")
    rodzaj = _rodzaj(sciezka)
    if rodzaj == "pdf":
        return _pdf(sciezka), "pdf"
    if rodzaj == "docx":
        return _docx(sciezka), "docx"
    if rodzaj == "odt":
        return _odt(sciezka), "odt"
    return sciezka.read_text(encoding="utf-8", errors="replace"), "txt"


if __name__ == "__main__":
    for arg in sys.argv[1:]:
        try:
            t, r = wczytaj(arg)
            print(f"{arg}  [{r}]  {len(t.split())} slow")
        except BladOdczytu as e:
            print(f"{arg}: {e}")
